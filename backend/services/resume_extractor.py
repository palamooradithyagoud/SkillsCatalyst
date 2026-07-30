import io
import re
import logging
import os
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# Configurable file size limit — defaults to 10 MB, overridable via env.
_MAX_MB = int(os.getenv("MAX_RESUME_MB", "10"))
MAX_FILE_SIZE_BYTES = _MAX_MB * 1024 * 1024


# ---------------------------------------------------------------------------
# Text Cleaning
# ---------------------------------------------------------------------------

def clean_extracted_text(text: str) -> str:
    """
    Normalizes and cleans text extracted from PDF, DOCX, or text files.

    - Removes control characters and null bytes.
    - Normalizes horizontal spaces while preserving line breaks.
    - Preserves bullet points (•, -, *, etc.), headings, URLs, emails, phones.
    - Collapses excessive blank lines (max 2 consecutive newlines).
    """
    if not text:
        return ""

    # Strip null bytes and non-printable control characters, keep \n \r \t
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

    # Normalize non-breaking and zero-width spaces
    text = text.replace('\xa0', ' ').replace('\u200b', '')

    # Trim each line individually
    lines = [line.strip() for line in text.splitlines()]
    cleaned = '\n'.join(lines)

    # Collapse multiple horizontal spaces to a single space (per line)
    cleaned = re.sub(r'[ \t]+', ' ', cleaned)

    # Collapse 3+ consecutive newlines to 2 (preserve paragraph spacing)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    return cleaned.strip()


# ---------------------------------------------------------------------------
# PDF Extraction
# ---------------------------------------------------------------------------

def _extract_pdf_pymupdf(content: bytes) -> str:
    """Extracts text from PDF bytes via PyMuPDF (fitz)."""
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as e:
        raise ValueError(f"Failed to open PDF document: {e}") from e

    if doc.is_encrypted:
        raise ValueError(
            "PDF is encrypted or password-protected. Please upload an unencrypted file."
        )

    pages: list[str] = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        page_text = page.get_text("text")
        if page_text and page_text.strip():
            pages.append(page_text.strip())

    doc.close()
    return "\n\n".join(pages)


def _extract_pdf_pdfplumber(content: bytes) -> str:
    """Fallback PDF extractor using pdfplumber (better for complex layouts)."""
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        logger.warning("pdfplumber not installed — skipping fallback extraction.")
        return ""

    pages: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages.append(page_text.strip())
    except Exception as e:
        logger.warning(f"pdfplumber fallback failed: {e}")
        return ""

    return "\n\n".join(pages)


def extract_text_from_pdf(content: bytes) -> str:
    """
    Extracts plain text from PDF bytes.

    Strategy:
      1. Try PyMuPDF (fast, handles most PDFs).
      2. Fall back to pdfplumber if PyMuPDF returns empty (image-heavy/complex layouts).
      3. Raise ValueError if both fail or document is encrypted / unreadable.
    """
    logger.debug("Attempting PDF extraction via PyMuPDF.")
    raw_text = _extract_pdf_pymupdf(content)  # raises on encryption

    if not raw_text.strip():
        logger.info("PyMuPDF returned empty text — trying pdfplumber fallback.")
        raw_text = _extract_pdf_pdfplumber(content)

    if not raw_text.strip():
        raise ValueError(
            "No readable text found in PDF. "
            "The document may be scanned or contain only images. "
            "Please upload a text-based PDF."
        )

    logger.info(f"PDF extraction successful ({len(raw_text)} chars before cleaning).")
    return raw_text


# ---------------------------------------------------------------------------
# DOCX Extraction
# ---------------------------------------------------------------------------

def extract_text_from_docx(content: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    Captures paragraphs and tables while preserving reading order.
    """
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise ValueError("python-docx is not installed. Cannot parse DOCX files.") from e

    try:
        stream = io.BytesIO(content)
        doc = docx.Document(stream)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {e}") from e

    lines: list[str] = []

    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            lines.append(paragraph.text.strip())

    # Include table cell text so skills/education tables aren't lost
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)

    raw_text = "\n".join(lines)

    if not raw_text.strip():
        raise ValueError("DOCX document is empty or contains no readable text.")

    logger.info(f"DOCX extraction successful ({len(raw_text)} chars before cleaning).")
    return raw_text


# ---------------------------------------------------------------------------
# Plain-text / Markdown Extraction
# ---------------------------------------------------------------------------

def extract_text_from_txt(content: bytes) -> str:
    """
    Decodes plain text or Markdown bytes.
    Tries UTF-8 first, falls back to Latin-1 / CP1252.
    """
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            decoded = content.decode(encoding)
            logger.info(
                f"TXT/MD decoded with {encoding} ({len(decoded)} chars before cleaning)."
            )
            return decoded
        except (UnicodeDecodeError, LookupError):
            continue

    raise ValueError(
        "Failed to decode text file. Ensure the file is saved as UTF-8."
    )


# ---------------------------------------------------------------------------
# Master Parse & Clean
# ---------------------------------------------------------------------------

def parse_and_clean_document(content: bytes, filename: str) -> str:
    """
    Validates, parses, and cleans a resume document.

    Supports: PDF, DOCX, DOC, TXT, MD.
    Returns clean plain text or raises ValueError with a clear message.
    """
    logger.info(f"Processing resume: '{filename}' ({len(content)} bytes).")

    if not content or len(content) == 0:
        raise ValueError("Uploaded file is empty (0 bytes).")

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(
            f"File size ({len(content) // (1024 * 1024)} MB) exceeds the "
            f"{_MAX_MB} MB limit. Please upload a smaller file."
        )

    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        raw_text = extract_text_from_pdf(content)
    elif filename_lower.endswith((".docx", ".doc")):
        raw_text = extract_text_from_docx(content)
    elif filename_lower.endswith((".txt", ".md")):
        raw_text = extract_text_from_txt(content)
    else:
        raise ValueError(
            "Unsupported file format. Please upload a PDF, DOCX, TXT, or MD file."
        )

    cleaned_text = clean_extracted_text(raw_text)

    if not cleaned_text or len(cleaned_text) < 10:
        raise ValueError(
            "Extracted text is too short to analyse. "
            "The file may be empty or contain only images."
        )

    logger.info(
        f"Resume '{filename}' processed successfully: {len(cleaned_text)} chars."
    )
    return cleaned_text
